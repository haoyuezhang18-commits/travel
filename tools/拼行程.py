#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
按路线编号，把模板里那张表**连同图片和排版原样搬进**新的交付文档，再批量替换变量。

不重打内容、不丢图片——这是做逐日行程的正确姿势。

用法：
    python3 tools/拼行程.py 配方.json 输出.docx

配方.json 格式：
{
  "标题": "西意法15天行程　文档1",
  "概述": ["✓ 9.09 周三 巴塞罗那：...", "..."],
  "全局替换": {"8月": "9月"},
  "天": [
    {"标题": "D1｜9.09（周三）巴塞罗那", "编号": "BCN-R02",
     "说明": "抵达日，上午落地可当完整一天用",
     "替换": {"5 月": "9 月", "09:30 - 12:00": "10:30 - 12:15"}}
  ]
}
"""
import os, sys, csv, json, copy, argparse, io, hashlib, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LOCATOR = os.path.join(HERE, 'catalog', '路线定位表.csv')
FONT, SIZE = 'HYJunHei-EEJ', Pt(9)
R_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'


def find_root(arg=None):
    if arg:
        return os.path.expanduser(arg)
    cfg = os.path.join(HERE, 'config', 'local-paths.yaml')
    if os.path.exists(cfg):
        for line in open(cfg, encoding='utf-8'):
            if line.strip().startswith('template_root:'):
                v = line.split(':', 1)[1].strip().strip('"\'')
                if v and not v.startswith('/path/to'):
                    return os.path.expanduser(v)
    return os.path.expanduser('~/文件夹/旅游通用模板')


def locate(code):
    with open(LOCATOR, encoding='utf-8') as f:
        rows = [r for r in csv.DictReader(f) if r['路线编号'].upper() == code.upper()]
    rows = [r for r in rows if not r.get('备注')] or rows
    if not rows:
        raise SystemExit(f'找不到路线编号 {code}')
    return rows[0]


def table_images(tbl):
    return len(list(tbl._tbl.iter(qn('a:blip'))))


def get_table(src_doc, idx):
    """定位表：伦敦模板首张是总索引表，需跳过。"""
    tables = src_doc.tables
    off = 1 if (tables and tables[0].rows and '路线编号' in tables[0].rows[0].cells[0].text) else 0
    return tables[idx - 1 + off]


def copy_images(tbl_el, src_part, dst_part, seen):
    """把表格里引用的图片搬到新文档，并重写 rId。

    不能直接 relate_to 源文档的 part：不同模板里都有 image1.png，
    partname 会在新包里撞车，产出重复 zip 条目、文件损坏。
    这里改为取出图片二进制，交给 get_or_add_image 重新登记，
    它会按 sha1 去重并分配唯一 partname。
    """
    n = 0
    for el in tbl_el.iter():
        for attr in (R_NS + 'embed', R_NS + 'link', R_NS + 'id'):
            rid = el.get(attr)
            if not rid:
                continue
            try:
                part = src_part.related_parts[rid]
            except KeyError:
                continue
            blob = part.blob
            key = hashlib.sha1(blob).hexdigest()
            if key not in seen:
                new_rid, _ = dst_part.get_or_add_image(io.BytesIO(blob))
                seen[key] = new_rid
            el.set(attr, seen[key])
            n += 1
    return n


TR = qn('w:tr'); TC = qn('w:tc'); WT = qn('w:t')


def row_text(tr, col=None):
    tcs = tr.findall(TC)
    if col is None:
        return ' '.join(''.join(t.text or '' for t in tc.iter(WT)) for tc in tcs)
    if col >= len(tcs):
        return ''
    return ''.join(t.text or '' for t in tcs[col].iter(WT))


def set_cell(tr, col, text):
    """把某格文字整体换掉，保留该格第一个 run 的格式；● 分点用真换行。"""
    tcs = tr.findall(TC)
    if col >= len(tcs):
        return
    tc = tcs[col]
    ps = tc.findall(qn('w:p'))
    keep = ps[0]
    for extra in ps[1:]:
        tc.remove(extra)
    runs = keep.findall(qn('w:r'))
    proto = copy.deepcopy(runs[0]) if runs else None
    for r in runs:
        keep.remove(r)
    for i, line in enumerate(str(text).split('\n')):
        if proto is not None:
            r = copy.deepcopy(proto)
            for t in r.findall(WT):
                r.remove(t)
            for br in r.findall(qn('w:br')):
                r.remove(br)
        else:
            r = keep.makeelement(qn('w:r'), {})
        if i:
            r.append(r.makeelement(qn('w:br'), {}))
        t = r.makeelement(WT, {})
        t.text = line
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        keep.append(r)


COLS = {'时间段': 0, '行程内容': 1, '交通': 2, '备注': 3}


def find_rows(rows, kw, whole=False, col=None, exact=False):
    """默认只在「行程内容」列匹配。

    备注列里经常顺带提到别的地名（例如古罗马广场的备注写着
    "从后门6出口出去就是真理之口"），整行匹配会改错行。
    """
    if exact:
        hit = [tr for tr in rows if row_text(tr, 1).strip() == kw]
        if hit:
            return hit
    if col is not None:
        return [tr for tr in rows if kw in row_text(tr, COLS.get(col, col))]
    if whole:
        return [tr for tr in rows if kw in row_text(tr)]
    hit = [tr for tr in rows if kw in row_text(tr, 1)]
    return hit if hit else [tr for tr in rows if kw in row_text(tr, 0)]


def apply_row_ops(tbl_el, day):
    """删 / 改 / 增 / 重排——字符串替换做不到的都在这里。"""
    log = []
    def data_rows():
        return tbl_el.findall(TR)[1:]

    for item in day.get('删', []):
        kw, col = (item, None) if isinstance(item, str) else (item['配'], item.get('配列'))
        hit = find_rows(data_rows(), kw, whole=day.get('整行匹配', False), col=col)
        if not hit:
            log.append(f'⚠ 删除失败，找不到「{kw}」')
        for tr in hit:
            tbl_el.remove(tr); log.append(f'删除「{kw}」')

    for op in day.get('改', []):
        kw = op['配']
        hit = find_rows(data_rows(), kw, whole=op.get('整行匹配', False),
                        col=op.get('配列'), exact=op.get('精确', False))
        if not hit:
            log.append(f'⚠ 修改失败，找不到「{kw}」'); continue
        if len(hit) > 1 and not op.get('允许多行'):
            log.append(f'⚠ 「{kw}」匹配到 {len(hit)} 行，只改第一行（如需全改加 允许多行:true）')
        for tr in (hit if op.get('允许多行') else hit[:1]):
            for ci, key in enumerate(['时间段', '行程内容', '交通', '备注']):
                if key in op:
                    set_cell(tr, ci, op[key])
            log.append(f'修改「{kw}」')

    for op in day.get('增', []):
        rows = tbl_el.findall(TR)
        proto = copy.deepcopy(rows[-1] if len(rows) > 1 else rows[0])
        for dr in proto.iter(qn('w:drawing')):     # 新行不带图
            dr.getparent().remove(dr)
        for ci, key in enumerate(['时间段', '行程内容', '交通', '备注']):
            set_cell(proto, ci, op.get(key, ''))
        anchor = op.get('位置', '尾')
        if anchor == '首':
            rows[0].addnext(proto)
        elif anchor == '尾':
            rows[-1].addnext(proto)
        else:
            hit = find_rows(tbl_el.findall(TR)[1:], anchor)
            if not hit:
                log.append(f'⚠ 插入位置找不到「{anchor}」，改放到表尾')
                tbl_el.findall(TR)[-1].addnext(proto)
            else:
                hit[-1].addnext(proto)
        log.append(f"新增「{str(op.get('行程内容',''))[:16]}」")

    if day.get('序'):
        rows = tbl_el.findall(TR)
        head, data = rows[0], rows[1:]
        ordered, rest = [], list(data)
        for kw in day['序']:
            hit = [tr for tr in rest if kw in row_text(tr, 1)] or [tr for tr in rest if kw in row_text(tr)]
            if hit:
                # 只取「行程内容」与首个命中完全相同的那一组：
                # 既能让交通备选这类重复行整组移动，又不会把
                # 「中央市场」误吞掉「中央市场 → 学院美术馆」。
                same = row_text(hit[0], 1).strip()
                group = [tr for tr in hit if row_text(tr, 1).strip() == same] or hit[:1]
                for tr in group:
                    ordered.append(tr); rest.remove(tr)
            else:
                log.append(f'⚠ 排序找不到「{kw}」')
        ordered += rest
        for tr in data:
            tbl_el.remove(tr)
        prev = head
        for tr in ordered:
            prev.addnext(tr); prev = tr
        log.append('已按指定顺序重排')
    return log


def replace_in_table(tbl_el, mapping):
    """在表格的文字节点上做替换，保留原有格式与图片。"""
    hits = 0
    for t in tbl_el.iter(qn('w:t')):
        if not t.text:
            continue
        new = t.text
        for a, b in mapping.items():
            if a in new:
                new = new.replace(a, b)
        if new != t.text:
            t.text = new
            hits += 1
    return hits


def styled(p, text, bold=False, size=SIZE, color=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = size
    r.bold = bold
    r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    if color:
        r.font.color.rgb = color
    return p


TIME_RE = re.compile(r'(\d{1,2})\s*[:：]\s*(\d{2})')


def audit(tbl_el, day):
    """生成后自检：时间倒流、旧客户信息残留、重复时段。"""
    warn = []
    rows = tbl_el.findall(TR)[1:]
    mins, labels = [], []
    for tr in rows:
        m = TIME_RE.search(row_text(tr, 0))
        if m:
            mins.append(int(m.group(1)) * 60 + int(m.group(2)))
            labels.append(row_text(tr, 0).strip()[:14] + ' ' + row_text(tr, 1).strip()[:16])
    for i in range(1, len(mins)):
        if mins[i] < mins[i - 1] - 5:
            warn.append(f'时间倒流：{labels[i-1]} → {labels[i]}')
    exact = {}
    for tr in rows:
        span = row_text(tr, 0).strip()
        body = row_text(tr, 1).strip()
        if not span or '备选' in body or '备选' in row_text(tr, 2):
            continue
        exact.setdefault(span, []).append(body[:20])
    for span, bodies in exact.items():
        if len(bodies) > 1:
            warn.append(f'时段完全重复（{span}）：' + ' / '.join(bodies))
    # 只查「行程内容」列：备注里顺带提一句别的景点是正常的
    body = ' '.join(row_text(tr, 1) for tr in rows)
    for bad in day.get('不应出现', []):
        if bad in body:
            warn.append(f'残留旧信息：「{bad}」')
    return warn


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('recipe'); ap.add_argument('out'); ap.add_argument('--root')
    a = ap.parse_args()
    cfg = json.load(open(a.recipe, encoding='utf-8'))
    root = find_root(a.root)

    doc = Document()
    st = doc.styles['Normal']
    st.font.name = FONT; st.font.size = SIZE
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    for s in doc.sections:
        s.left_margin = s.right_margin = Pt(36)

    styled(doc.add_paragraph(), cfg.get('标题', ''), bold=True, size=Pt(16))
    for line in cfg.get('概述', []):
        styled(doc.add_paragraph(), line)
    doc.add_paragraph()

    total_img = 0
    cache = {}
    seen_imgs = {}
    problems = []
    for day in cfg['天']:
        sources = day.get('取') or [{'编号': day['编号']}]

        styled(doc.add_paragraph(), day['标题'], bold=True, size=Pt(13))
        if day.get('说明'):
            styled(doc.add_paragraph(), day['说明'], size=Pt(8.5),
                   color=RGBColor(0x88, 0x44, 0x00))

        new_el, n, srcnote = None, 0, []
        for si, sc in enumerate(sources):
            row = locate(sc['编号'])
            path = os.path.join(root, row['模板文件（相对旅游通用模板/）'])
            if path not in cache:
                cache[path] = Document(path)
            src = cache[path]
            tbl = get_table(src, int(row['文件内第几张表']))
            srcnote.append(f"{sc['编号']}(图{table_images(tbl)})")
            part = copy.deepcopy(tbl._tbl)
            n += copy_images(part, src.part, doc.part, seen_imgs)
            keep = sc.get('要')
            if keep is not None:                       # 只取指定的行
                for tr in part.findall(TR)[1:]:
                    if not any(k in row_text(tr, 1) or k in row_text(tr, 0) for k in keep):
                        part.remove(tr)
            if new_el is None:
                new_el = part
            else:                                       # 追加到主表后面
                if keep is None:
                    rows_to_add = part.findall(TR)[1:]
                else:
                    rows_to_add = part.findall(TR)[1:]
                for tr in rows_to_add:
                    new_el.append(tr)
        code = ' + '.join(srcnote)
        total_img += n
        mapping = dict(cfg.get('全局替换', {})); mapping.update(day.get('替换', {}))
        hits = replace_in_table(new_el, mapping)
        oplog = apply_row_ops(new_el, day)
        # 必须插在 sectPr 之前、且紧跟本天标题；
        # 直接 body.append 会把所有表格堆到文档末尾、彼此黏连。
        anchor_p = doc.add_paragraph()
        anchor_p._p.addprevious(new_el)
        gap = anchor_p
        gap.paragraph_format.space_after = Pt(16)
        gap.paragraph_format.space_before = Pt(8)
        styled(gap, '')
        print(f"  {day['标题'][:20]:22} ← {code} 图{n} 替换{hits}处 " +
              (f"行操作{len(oplog)}项" if oplog else ""))
        for L in oplog:
            if L.startswith('⚠'):
                print('      ' + L); problems.append(f"{day['标题'][:12]}｜{L}")
        for w in audit(new_el, day):
            print('      ⚠ ' + w); problems.append(f"{day['标题'][:12]}｜{w}")

    doc.save(a.out)
    print(f'\n已生成：{a.out}　（共 {len(cfg["天"])} 天，搬运图片 {total_img} 处）')
    if problems:
        print(f'\n❌ 自检发现 {len(problems)} 个问题，必须修完再交付：')
        for p in problems:
            print('   ' + p)
    else:
        print('\n✅ 自检通过：时间顺序正常、无重复时段、无旧客户信息残留')


if __name__ == '__main__':
    main()
