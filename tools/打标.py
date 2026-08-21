# -*- coding: utf-8 -*-
"""给本地模板 DOCX 打上路线编号标签，使 AI 和人都能直接定位到某一张表，无需通读全文。"""
import os,sys,copy
sys.path.insert(0,os.path.dirname(os.path.abspath(__file__)))
from mapping import MAP
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT=os.path.expanduser('~/文件夹/旅游通用模板/')

def mk_para(doc, text, bold=True, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p

def prev_text(tbl):
    el = tbl._tbl.getprevious()
    hops = 0
    while el is not None and hops < 3:
        t = ''.join(el.itertext()) if el is not None else ''
        if t.strip(): return t
        el = el.getprevious(); hops += 1
    return ''

changed = []
for rel, codes in MAP.items():
    path = ROOT + rel
    doc = Document(path)
    tables = doc.tables
    touched = False

    # 1) 每张表前插入编号标题（已有则跳过）
    for tbl, (code, name) in zip(tables, codes):
        if code == '__INDEX__':  # 伦敦那张总索引表本身
            continue
        if code in prev_text(tbl):
            continue
        p = mk_para(doc, f'【{code}】{name}')
        tbl._tbl.addprevious(p._p)
        touched = True

    # 2) 多版本文件：顶部加一张路线索引（伦敦已自带，跳过）
    real = [(c, n) for c, n in codes if c != '__INDEX__']
    has_index_tbl = any(c == '__INDEX__' for c, n in codes)
    if len(real) > 1 and not has_index_tbl:
        body = doc.element.body
        head = mk_para(doc, f'本文件路线索引（共 {len(real)} 条，Ctrl+F 搜编号直接跳转）', bold=True, size=12)
        body.insert(0, head._p)
        for i, (c, n) in enumerate(real):
            line = mk_para(doc, f'　【{c}】{n}', bold=False, size=10.5)
            body.insert(1 + i, line._p)
        touched = True

    if touched:
        doc.save(path)
        changed.append((rel, len(real)))

print(f'已处理 {len(changed)} 个文件：')
for rel, n in changed:
    print(f'  {n:2}条  {rel.split("/")[-1]}')
